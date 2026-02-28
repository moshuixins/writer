import json
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from app.config import get_settings
from app.errors import DocxGenerationError, logger

settings = get_settings()


class DocxGenerator:
    """公文docx生成引擎，遵循GB/T 9704-2012标准"""

    # 字体常量
    TITLE_FONT = "方正小标宋简体"
    TITLE_FALLBACK = "SimSun"
    BODY_FONT = "仿宋_GB2312"
    BODY_FALLBACK = "FangSong"
    H1_FONT = "黑体"
    H2_FONT = "楷体_GB2312"
    H2_FALLBACK = "KaiTi"

    def _set_font(self, run, font_name: str, fallback: str, size_pt: int):
        """设置字体，含中文回退"""
        run.font.size = Pt(size_pt)
        run.font.name = font_name
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _setup_page(self, doc: Document):
        """设置页面布局"""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)

    def _set_line_spacing(self, paragraph, pt_value: int = 28):
        """设置固定行距"""
        from docx.oxml import OxmlElement
        pPr = paragraph._element.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), str(pt_value * 20))
        spacing.set(qn("w:lineRule"), "exact")
        pPr.append(spacing)

    def _add_title(self, doc: Document, title: str):
        """添加公文标题"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        self._set_font(run, self.TITLE_FONT, self.TITLE_FALLBACK, 22)
        run.bold = True
        self._set_line_spacing(p, 28)

    def _add_body_paragraph(self, doc: Document, text: str):
        """添加正文段落"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)  # 2字符缩进
        run = p.add_run(text)
        self._set_font(run, self.BODY_FONT, self.BODY_FALLBACK, 16)
        self._set_line_spacing(p, 28)

    def _add_heading(self, doc: Document, text: str, level: int = 1):
        """添加标题（一级用黑体，二级用楷体）"""
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        if level == 1:
            self._set_font(run, self.H1_FONT, self.H1_FONT, 16)
            run.bold = True
        else:
            self._set_font(run, self.H2_FONT, self.H2_FALLBACK, 16)
        self._set_line_spacing(p, 28)

    def _add_signature(self, doc: Document, org: str, date: str):
        """添加落款"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(org)
        self._set_font(run, self.BODY_FONT, self.BODY_FALLBACK, 16)

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = p2.add_run(date)
        self._set_font(run2, self.BODY_FONT, self.BODY_FALLBACK, 16)

    def generate(self, content_json: dict) -> str:
        """根据结构化JSON生成公文docx，返回文件路径"""
        if not isinstance(content_json, dict):
            raise DocxGenerationError(detail="content_json 必须是字典")
        if not content_json.get("title") and not content_json.get("body_sections"):
            raise DocxGenerationError(detail="至少需要 title 或 body_sections")

        try:
            doc = Document()
            self._setup_page(doc)

            # 标题
            title = content_json.get("title", "")
            if title:
                self._add_title(doc, title)

            # 主送机关
            recipients = content_json.get("recipients", "")
            if recipients:
                self._add_body_paragraph(doc, recipients + "：")

            # 正文各段
            for section in content_json.get("body_sections", []):
                heading = section.get("heading", "")
                content = section.get("content", "")
                level = section.get("level", 0)

                if heading:
                    self._add_heading(doc, heading, level=level)
                if content:
                    self._add_body_paragraph(doc, content)

            # 落款
            signing_org = content_json.get("signing_org", "")
            date = content_json.get("date", "")
            if signing_org:
                self._add_signature(doc, signing_org, date)

            # 保存文件
            export_dir = Path(settings.export_dir)
            export_dir.mkdir(parents=True, exist_ok=True)
            filename = f"{uuid.uuid4().hex}.docx"
            filepath = export_dir / filename
            doc.save(str(filepath))
            logger.info("Docx generated: %s", filepath)
            return str(filepath)
        except DocxGenerationError:
            raise
        except Exception as e:
            logger.error("Docx generation failed: %s", e)
            raise DocxGenerationError(detail=str(e))
